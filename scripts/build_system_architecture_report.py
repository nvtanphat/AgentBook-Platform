from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Bao_cao_kien_truc_he_thong_AgentBook.docx"

BLACK = RGBColor(0, 0, 0)
BORDER = "DADCE0"
LIGHT_GRAY = "F3F4F6"
MID_GRAY = "E5E7EB"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_paragraph_black(paragraph):
    for run in paragraph.runs:
        run.font.color.rgb = BLACK


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = BLACK
        p.add_run(text[len(bold_prefix):]).font.color.rgb = BLACK
    else:
        p.add_run(text).font.color.rgb = BLACK
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    r = p.add_run(text)
    r.font.color.rgb = BLACK
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.color.rgb = BLACK
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = BLACK
        run.bold = True
    return p


def add_feature_box(doc, title, reason):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, color="C7CBD1", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Tính năng nổi bật: ")
    r.bold = True
    r.font.color.rgb = BLACK
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = BLACK
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run("Vì sao hay: ")
    r.bold = True
    r.font.color.rgb = BLACK
    r = p2.add_run(reason)
    r.font.color.rgb = BLACK
    doc.add_paragraph()


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "Thành phần"
    hdr[1].text = "Vai trò trong hệ thống"
    for cell in hdr:
        set_cell_shading(cell, MID_GRAY)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = BLACK
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    set_table_width(table, [2.1, 4.4])
    set_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    run.font.color.rgb = BLACK
    doc.add_paragraph()


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after in [
        ("Heading 1", 20, 20, 6),
        ("Heading 2", 16, 18, 6),
        ("Heading 3", 14, 16, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.font.color.rgb = BLACK
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def build():
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Báo cáo tổng quan kiến trúc hệ thống AgentBook")
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.color.rgb = BLACK
    run.bold = False

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(14)
    r = meta.add_run("Ngày lập: 04/06/2026 | Phạm vi: Backend, AI/RAG, Frontend, triển khai và vận hành")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK

    add_heading(doc, "1. Tổng quan Kiến trúc Hệ thống", 1)
    add_para(
        doc,
        "AgentBook là trợ lý học tập dựa trên bằng chứng, biến tài liệu học tập nhiều định dạng thành một bề mặt tri thức có thể hỏi đáp, trích dẫn và kiểm chứng. Kiến trúc được chia thành năm lớp chính: giao diện React, API FastAPI, pipeline xử lý tài liệu, lớp RAG/agentic reasoning và các kho lưu trữ MongoDB/Qdrant/Redis.",
    )
    add_para(
        doc,
        "Luồng tổng quát: người dùng tải tài liệu lên; backend kiểm tra an toàn, phân tích tài liệu, OCR hoặc nhận dạng âm thanh khi cần; dữ liệu được chuẩn hóa thành block/chunk có tọa độ bằng chứng; embedding được ghi vào Qdrant; metadata, chunk, entity và quan hệ được giữ trong MongoDB; truy vấn người dùng đi qua phân loại ý định, truy hồi lai, rerank, tổng hợp câu trả lời và kiểm tra groundedness trước khi trả về UI.",
    )
    add_feature_box(
        doc,
        "Trích dẫn bám tới block/page/bbox/audio timestamp",
        "Người dùng không chỉ thấy câu trả lời mà còn có thể kiểm tra câu trả lời dựa trên đúng trang, vùng văn bản, hàng bảng hoặc đoạn âm thanh nguồn. Đây là điểm quan trọng để giảm ảo giác và phù hợp với môi trường học thuật.",
    )

    add_kv_table(
        doc,
        [
            ("Frontend", "React + Vite hiển thị workspace, chat, nguồn tài liệu, bằng chứng, graph/mindmap và agent trace."),
            ("API backend", "FastAPI cung cấp các route /api/v1 cho query, materials, collections, graph, auth, admin và evaluation."),
            ("Xử lý tài liệu", "Docling, OCR, parser bảng, audio parser, chunking, enrichment và entity/relation extraction."),
            ("RAG & agentic layer", "Hybrid retrieval, cross-encoder rerank, planner/director/critic/synthesizer/guardrails/SLEC."),
            ("Lưu trữ", "MongoDB/Beanie cho metadata và evidence; Qdrant cho vector; Redis cho queue/cache; file gốc trong data volume."),
        ],
    )

    add_heading(doc, "2. Công nghệ phía Máy chủ (Backend & Database)", 1)
    add_bullet(doc, "FastAPI 0.115.5 và Uvicorn là lớp API chính, có health check, rate limiting và cấu trúc route theo /api/v1.")
    add_bullet(doc, "Pydantic v2 và pydantic-settings chuẩn hóa schema, cấu hình và hợp đồng dữ liệu giữa service, API và UI.")
    add_bullet(doc, "MongoDB được truy cập qua Motor/Beanie để lưu collection, material, chunk, block, entity, relation, chat memory và feedback.")
    add_bullet(doc, "Qdrant lưu dense/sparse vector để phục vụ hybrid retrieval; Redis dùng cho Celery broker/result backend và cache.")
    add_bullet(doc, "Celery worker tách ingestion nặng khỏi API request, tránh làm nghẽn trải nghiệm upload và truy vấn.")
    add_feature_box(
        doc,
        "Tách API và worker ingestion",
        "Các tác vụ parse/OCR/embedding có độ trễ cao được đưa sang hàng đợi, giúp API vẫn phản hồi ổn định và có thể mở rộng worker độc lập khi khối lượng tài liệu tăng.",
    )

    add_heading(doc, "3. Khối Công nghệ Trí tuệ Nhân tạo (AI & RAG Pipeline)", 1)
    add_para(doc, "Khối AI là phần lõi của hệ thống, gồm pipeline nhập liệu và pipeline trả lời truy vấn.")
    add_heading(doc, "3.1 Pipeline nhập liệu", 2)
    for item in [
        "Docling phân tích PDF/DOCX/PPTX thành block văn bản, bảng, hình, công thức.",
        "EasyOCR và VLM fallback xử lý trang scan hoặc chữ viết tay; Faster-Whisper xử lý audio có timestamp.",
        "Layout normalizer, reading order, table parser và chunking token-aware giữ ngữ cảnh và nguồn gốc từng đoạn.",
        "Contextual enrichment thêm tóm tắt lân cận; entity/relation extraction tạo lớp knowledge graph.",
        "BGE-M3 tạo dense vector 1024 chiều và sparse vector trong một lượt, sau đó ghi vào Qdrant.",
    ]:
        add_number(doc, item)
    add_heading(doc, "3.2 Pipeline truy vấn", 2)
    for item in [
        "IntentClassifier phát hiện chitchat/off-topic để từ chối nhanh thay vì chạy toàn bộ RAG.",
        "QueryProcessor xử lý ngôn ngữ, anaphora, truy vấn song ngữ và multi-query khi cần tăng recall.",
        "PlannerAgent và RetrieverDirector chọn chiến lược: text retrieval, per-source search hoặc graph traversal.",
        "Hybrid retrieval dùng dense + sparse, hợp nhất bằng Reciprocal Rank Fusion rồi cross-encoder rerank top candidates.",
        "SynthesizerAgent sinh câu trả lời có citation; Guardrails, NLI verifier, contradiction detector và Sentence-Level Coverage Gate loại hoặc hedge câu thiếu bằng chứng.",
    ]:
        add_number(doc, item)
    add_feature_box(
        doc,
        "Agentic RAG có trạng thái và trace",
        "Thay vì để LLM tự suy luận tự do, hệ thống chia truy vấn thành các bước có hợp đồng rõ: plan, retrieve, critique, rerank, synthesize, verify. UI có thể hiển thị trace để người dùng biết câu trả lời được tạo ra như thế nào.",
    )
    add_feature_box(
        doc,
        "Cross-lingual VI-EN retrieval",
        "Sinh viên có thể hỏi tiếng Việt trên giáo trình tiếng Anh hoặc ngược lại. Truy vấn gốc giữ recall theo ngôn ngữ nguồn, bản dịch bắt paraphrase, RRF hợp nhất kết quả nên ít mất bằng chứng.",
    )

    add_heading(doc, "4. Công nghệ Giao diện Người dùng (Frontend)", 1)
    add_bullet(doc, "React 18, TypeScript và Vite 6 là nền tảng phát triển UI.")
    add_bullet(doc, "React Router quản lý luồng login/workspace; state được tách trong auth/theme/workspace context.")
    add_bullet(doc, "Các component chính gồm ChatPanel, SourcesPanel, EvidencePanel, GraphCanvas, ReasoningTrace, DebugModal và SettingsModal.")
    add_bullet(doc, "ReactFlow và Dagre hỗ trợ graph/mindmap; lucide-react dùng cho icon; Tailwind/PostCSS xử lý style.")
    add_bullet(doc, "Client API TypeScript gom hợp đồng gọi backend, giúp UI không gọi endpoint rải rác.")
    add_feature_box(
        doc,
        "Evidence-first UI",
        "Các citation, snippet, vùng highlight, audio player, graph và reasoning trace được đưa vào cùng workspace. Điều này biến hệ thống từ chatbot thành công cụ học tập có khả năng kiểm chứng.",
    )

    add_heading(doc, "5. Môi trường Triển khai & Vận hành (Deployment)", 1)
    add_para(doc, "Dự án có docker-compose cho API, worker, Qdrant và Redis. API và worker cùng build từ backend/Dockerfile nhưng khác command; Qdrant và Redis có healthcheck; volume data lưu tài liệu, vector DB và cấu hình runtime.")
    add_bullet(doc, "API expose 127.0.0.1:8000, Qdrant expose 6333/6334, Redis expose 6379.")
    add_bullet(doc, "Ollama được gọi qua host.docker.internal:11434, model local mặc định trong compose là qwen3:4b.")
    add_bullet(doc, "Cấu hình chia thành các file YAML: retrieval_config, guardrails_config, extraction_config, model_config, viz_config và logging_config.")
    add_bullet(doc, "Frontend build bằng npm script build; dev mode dùng Vite.")
    add_feature_box(
        doc,
        "Cấu hình theo lớp và healthcheck rõ ràng",
        "Việc tách retrieval, guardrails, extraction và model config giúp thử nghiệm nhanh mà không sửa code. Healthcheck giúp compose khởi động theo phụ thuộc thực tế thay vì chỉ theo thứ tự container.",
    )

    add_heading(doc, "6. Đánh giá Ưu điểm & Hạn chế", 1)
    add_heading(doc, "6.1 Ưu điểm", 2)
    for item in [
        "Groundedness mạnh: citation có block/page/bbox/timestamp và SLEC kiểm tra từng câu.",
        "Hỗ trợ dữ liệu học tập đa định dạng: PDF, slide, DOCX, bảng, ảnh scan, chữ viết tay và audio.",
        "Hybrid retrieval + rerank phù hợp câu hỏi học thuật vì kết hợp recall rộng và precision sau rerank.",
        "Kiến trúc agentic có trace, dễ debug hơn pipeline RAG hộp đen.",
        "Triển khai cục bộ tốt: dùng Ollama/local LLM, Qdrant và Redis qua Docker, giảm phụ thuộc dịch vụ bên ngoài.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "6.2 Hạn chế", 2)
    for item in [
        "Pipeline nhiều tầng nên độ trễ truy vấn có thể cao, đặc biệt khi chạy local LLM hoặc OCR trên CPU.",
        "Chất lượng OCR/chữ viết tay phụ thuộc ảnh đầu vào; dữ liệu nhiễu vẫn cần quality gate và kiểm tra thủ công.",
        "MongoDB không có trong docker-compose hiện tại, vì vậy môi trường triển khai cần cung cấp MongoDB bên ngoài hoặc bổ sung service.",
        "Nhiều model AI/embedding/rerank làm chi phí RAM/CPU/GPU tăng; cần profile hiệu năng trước khi triển khai nhiều người dùng.",
        "Guardrails giảm ảo giác nhưng không thay thế đánh giá học thuật; các câu trả lời quan trọng vẫn cần người dùng kiểm chứng bằng citation.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. Kết luận", 1)
    add_para(
        doc,
        "AgentBook có kiến trúc phù hợp cho trợ lý học tập dựa trên bằng chứng: backend rõ lớp, RAG pipeline giàu kiểm chứng, UI tập trung vào evidence và triển khai cục bộ tương đối gọn. Điểm hay nhất là hệ thống không chỉ trả lời mà còn cung cấp đường kiểm chứng từ câu trả lời về tài liệu nguồn. Hạn chế chính nằm ở độ phức tạp vận hành, tài nguyên tính toán và chất lượng đầu vào đa phương thức.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())

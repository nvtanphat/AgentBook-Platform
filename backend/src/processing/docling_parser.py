from __future__ import annotations

import logging
import os
import re
from pathlib import Path
from typing import Any
from uuid import NAMESPACE_URL, uuid5

from src.processing.types import BBox, BlockType, DependencyUnavailableError, ParsedBlock, ParsedDocument, ParsedPage

SUPPORTED_DOCLING_EXTENSIONS = {"pdf", "pptx", "docx"}
logger = logging.getLogger(__name__)
_BULLET_PREFIX = re.compile(r"^[\u25a1\u2022\u25cf\u25aa\u25ab\u2013\u2014\-\*\u25c6\u25c7\u25cb]\s")
_SECTION_SPLIT = re.compile(r"(?=(?:^|\n)(?:\d+(?:\.\d+)*\.?\s+|[A-Z][A-Za-z ]{2,}:))")


def _workspace_cache_dir(name: str) -> Path:
    return Path(__file__).resolve().parents[3] / "data" / "cache" / name


def _is_empty_figure_only(blocks: list[ParsedBlock]) -> bool:
    """True when the page's only blocks are figure placeholders with no text.

    Docling marks scanned-text legal PDF pages as a single FIGURE block with
    `needs_captioning=True` and empty content. Those pages must still go
    through EasyOCR — otherwise the entire page text is lost downstream.
    """
    if not blocks:
        return True
    return all(
        block.block_type == BlockType.FIGURE.value and not (block.content or "").strip()
        for block in blocks
    )


def _configure_docling_cache() -> None:
    os.environ.setdefault("USE_TF", "0")
    os.environ.setdefault("TRANSFORMERS_NO_TF", "1")
    os.environ.setdefault("USE_FLAX", "0")
    os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS_WARNING", "1")
    cache_dirs = {
        "HF_HOME": _workspace_cache_dir("huggingface"),
        "MODELSCOPE_CACHE": _workspace_cache_dir("modelscope"),
    }
    for name, path in cache_dirs.items():
        path.mkdir(parents=True, exist_ok=True)
        os.environ.setdefault(name, str(path))
    try:
        import huggingface_hub.file_download as hf_file_download
    except ImportError:
        return
    hf_file_download.are_symlinks_supported = lambda cache_dir=None: False


def _patch_docling_transformers_compat() -> None:
    try:
        import transformers
        import transformers.image_utils as image_utils
    except ImportError:
        return
    if not hasattr(image_utils, "VideoInput"):
        setattr(image_utils, "VideoInput", getattr(image_utils, "ImageInput", object))
    if not hasattr(image_utils, "VideoMetadata"):
        setattr(image_utils, "VideoMetadata", dict)
    if hasattr(transformers, "AutoModelForImageTextToText"):
        return
    fallback = getattr(transformers, "AutoModelForVision2Seq", None)
    if fallback is not None:
        setattr(transformers, "AutoModelForImageTextToText", fallback)


class DoclingParser:
    def parse(self, file_path: Path, *, language: str = "unknown") -> ParsedDocument:
        extension = file_path.suffix.lower().lstrip(".")
        if extension not in SUPPORTED_DOCLING_EXTENSIONS:
            raise ValueError(f"DoclingParser does not support .{extension}")

        try:
            self._ensure_docling_available()
        except ImportError as exc:
            if extension == "pdf":
                logger.warning("Docling is unavailable for PDF; using pypdf/OCR fallback", extra={"error": str(exc)})
                return self._pdf_text_fallback_document(file_path, language=language, parser_error=exc)
            raise DependencyUnavailableError("docling is required for PDF/PPTX/DOCX parsing") from exc

        converter = self._converter(extension)
        try:
            result = converter.convert(str(file_path))
        except Exception as exc:
            if extension == "pdf":
                logger.warning("Docling PDF parse failed, using pypdf text fallback", extra={"error": str(exc)})
                return self._pdf_text_fallback_document(file_path, language=language, parser_error=exc)
            raise

        pages = self._pages_from_export(result.document, file_path=file_path, extension=extension, language=language)
        extra = {"parser": "docling"}
        if extension == "pdf":
            extra["pdf_strategy"] = "docling_layout_first_text_ocr_missing_pages"
        return ParsedDocument(
            source_path=str(file_path),
            file_type=extension,
            language=language,
            pages=pages,
            extra=extra,
        )

    @staticmethod
    def _ensure_docling_available() -> None:
        _configure_docling_cache()
        _patch_docling_transformers_compat()
        from docling.document_converter import DocumentConverter  # noqa: F401

    @staticmethod
    def _converter(extension: str):
        from docling.document_converter import DocumentConverter

        if extension != "pdf":
            return DocumentConverter()

        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.document_converter import PdfFormatOption

        from src.core.config import get_settings as _get_settings
        _cfg = _get_settings()
        pdf_options = PdfPipelineOptions(
            do_ocr=False,
            ocr_batch_size=1,
            layout_batch_size=1,
            table_batch_size=1,
            queue_max_size=_cfg.docling_queue_max_size,
            images_scale=_cfg.docling_images_scale,
            # NOTE: do NOT set generate_picture_images=True here — retaining all
            # picture pixels across the document triggers std::bad_alloc on
            # image-heavy PDFs. Figure images are instead extracted lazily,
            # page-by-page, with PyMuPDF in the captioning step (memory-light).
        )
        return DocumentConverter(format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_options)})

    def _pages_from_export(self, document: Any, *, file_path: Path, extension: str, language: str) -> list[ParsedPage]:
        exported = self._export_dict(document)
        blocks = self._blocks_from_docling_dict(exported, language=language)
        if extension == "docx":
            blocks.extend(self._docx_table_blocks(file_path=file_path, language=language, reading_order_offset=len(blocks)))
        if not blocks:
            markdown = self._export_markdown(document)
            if markdown.strip():
                blocks = [
                    ParsedBlock(
                        block_id=self._stable_block_id(file_path, 1, 0, markdown[:80]),
                        block_index=0,
                        block_type=BlockType.PARAGRAPH.value,
                        content=markdown.strip(),
                        page_number=1,
                        language=language,
                        reading_order=0,
                        source="docling_markdown",
                    )
                ]
        pages_by_number: dict[int, list[ParsedBlock]] = {}
        for block in blocks:
            pages_by_number.setdefault(block.page_number, []).append(block)
        if extension == "pdf":
            self._add_pdf_text_fallback_pages(pages_by_number, file_path=file_path, language=language)
            self._add_easyocr_pages(pages_by_number, file_path=file_path, language=language)
        if not pages_by_number:
            return [ParsedPage(page_number=1, blocks=[])]
        return [
            ParsedPage(page_number=page_number, blocks=sorted(page_blocks, key=lambda block: block.reading_order))
            for page_number, page_blocks in sorted(pages_by_number.items())
        ]

    def _pdf_text_fallback_document(self, file_path: Path, *, language: str, parser_error: Exception) -> ParsedDocument:
        pages_by_number: dict[int, list[ParsedBlock]] = {}
        self._add_pdf_text_fallback_pages(pages_by_number, file_path=file_path, language=language)
        self._add_easyocr_pages(pages_by_number, file_path=file_path, language=language)
        pages = [
            ParsedPage(page_number=page_number, blocks=sorted(page_blocks, key=lambda block: block.reading_order))
            for page_number, page_blocks in sorted(pages_by_number.items())
        ] or [ParsedPage(page_number=1, blocks=[])]
        return ParsedDocument(
            source_path=str(file_path),
            file_type="pdf",
            language=language,
            pages=pages,
            extra={"parser": "pypdf_fallback", "docling_error": f"{type(parser_error).__name__}: {parser_error}"},
        )

    def _add_pdf_text_fallback_pages(
        self,
        pages_by_number: dict[int, list[ParsedBlock]],
        *,
        file_path: Path,
        language: str,
    ) -> None:
        try:
            from pypdf import PdfReader
        except ImportError:
            return

        try:
            reader = PdfReader(str(file_path))
        except Exception:
            return

        for page_index, page in enumerate(reader.pages, start=1):
            if page_index in pages_by_number:
                continue
            try:
                text = (page.extract_text() or "").strip()
            except Exception:
                continue
            if not text:
                continue
            pages_by_number[page_index] = self._blocks_from_plain_text_page(
                text,
                file_path=file_path,
                page_number=page_index,
                language=language,
                source="pypdf_text_fallback",
                fallback_reason="docling_missing_page",
            )

    def _blocks_from_plain_text_page(
        self,
        text: str,
        *,
        file_path: Path,
        page_number: int,
        language: str,
        source: str,
        fallback_reason: str,
    ) -> list[ParsedBlock]:
        blocks: list[ParsedBlock] = []
        for unit in self._split_plain_text_page(text):
            cleaned = unit.strip()
            if not cleaned:
                continue
            blocks.append(
                ParsedBlock(
                    block_id=self._stable_block_id(file_path, page_number, len(blocks), cleaned[:80]),
                    block_index=len(blocks),
                    block_type=self._classify_plain_text_block(cleaned),
                    content=cleaned,
                    page_number=page_number,
                    language=language,
                    reading_order=len(blocks),
                    source=source,
                    extra={"fallback_reason": fallback_reason},
                )
            )
        return blocks

    @staticmethod
    def _split_plain_text_page(text: str) -> list[str]:
        lines = [" ".join(line.split()) for line in text.splitlines() if line.strip()]
        if len(lines) <= 1:
            parts = [part.strip() for part in _SECTION_SPLIT.split(text) if part.strip()]
            return parts or [text.strip()]

        units: list[str] = []
        current: list[str] = []
        for line in lines:
            starts_new = (
                _BULLET_PREFIX.match(line) is not None
                or re.match(r"^\d+(?:\.\d+)*\.?\s+", line) is not None
                or (len(line) <= 90 and not line.endswith((".", ",", ";")))
            )
            current_is_heading = len(current) == 1 and len(current[0]) <= 90 and not current[0].endswith((".", ",", ";"))
            if current and (starts_new or current_is_heading):
                units.append(" ".join(current))
                current = []
            current.append(line)
        if current:
            units.append(" ".join(current))
        return units

    @staticmethod
    def _classify_plain_text_block(text: str) -> str:
        if "|" in text and "\n" in text:
            return BlockType.TABLE.value
        if _BULLET_PREFIX.match(text) or text.startswith(("-", "*", "\u2022")):
            return BlockType.LIST.value
        if len(text) <= 120 and not text.endswith((".", ";", ":")):
            return BlockType.HEADING.value
        return BlockType.PARAGRAPH.value

    def _add_easyocr_pages(
        self,
        pages_by_number: dict[int, list[ParsedBlock]],
        *,
        file_path: Path,
        language: str,
    ) -> None:
        if not file_path.exists():
            return
        try:
            import pypdfium2 as pdfium
        except ImportError as exc:
            raise DependencyUnavailableError("pypdfium2 is required to render scanned PDF pages for OCR") from exc

        from src.core.config import get_settings
        from src.processing.ocr_engine import EasyOCREngine, VietOCRRecognizer

        settings = get_settings()
        ocr_language = "vi" if language == "vi" else "en"
        output_dir = _workspace_cache_dir("pdf_page_images")
        output_dir.mkdir(parents=True, exist_ok=True)
        # Best-of-breed for Vietnamese scanned pages: EasyOCR detection + VietOCR
        # recognition (tone-accurate), honouring `ocr_recognition_engine`. The
        # dedicated image pipeline already does this; without it, EasyOCR reads
        # Vietnamese tone marks poorly ("hợp đồng" → "ợông đồng") on scanned PDFs.
        recognizer = None
        if ocr_language == "vi" and settings.ocr_recognition_engine == "vietocr":
            recognizer = VietOCRRecognizer(
                device=settings.ocr_vietocr_device,
                model_name=settings.ocr_vietocr_model_name,
            )
        engine = EasyOCREngine(lang=ocr_language, recognizer=recognizer)
        render_scale = settings.pdf_render_scale

        pdf = pdfium.PdfDocument(str(file_path))
        try:
            for page_index in range(len(pdf)):
                page_number = page_index + 1
                # Skip pages Docling already parsed with REAL text content.
                # Pages where the only blocks are empty figure placeholders
                # (typical for scanned legal PDFs Docling treats as images)
                # must still go through EasyOCR — otherwise the page text is lost.
                existing = pages_by_number.get(page_number, [])
                if existing and not _is_empty_figure_only(existing):
                    continue
                page = pdf[page_index]
                try:
                    bitmap = page.render(scale=render_scale)
                    image = bitmap.to_pil()
                    image_path = output_dir / f"{uuid5(NAMESPACE_URL, f'{file_path}:{page_number}').hex}.png"
                    image.save(image_path)
                finally:
                    page.close()

                parsed = engine.parse_image(image_path, language=language if language in {"vi", "en"} else "unknown")
                blocks = [
                    block.model_copy(update={"page_number": page_number, "block_index": index, "reading_order": index})
                    for index, block in enumerate(parsed.blocks)
                ]
                if blocks:
                    # Replace empty figure placeholder with the OCR'd text blocks.
                    pages_by_number[page_number] = blocks
        finally:
            pdf.close()

    @staticmethod
    def _export_dict(document: Any) -> dict[str, Any]:
        for method_name in ("export_to_dict", "model_dump", "dict"):
            method = getattr(document, method_name, None)
            if callable(method):
                data = method()
                if isinstance(data, dict):
                    return data
        return {}

    @staticmethod
    def _export_markdown(document: Any) -> str:
        method = getattr(document, "export_to_markdown", None)
        if callable(method):
            return str(method())
        return ""

    def _blocks_from_docling_dict(self, data: dict[str, Any], *, language: str) -> list[ParsedBlock]:
        candidates = self._collect_text_nodes(data)
        figure_nodes = self._collect_figure_nodes(data)
        blocks: list[ParsedBlock] = []

        for index, node in enumerate(candidates):
            text = self._node_text(node)
            if not text:
                continue
            label = str(node.get("label") or node.get("type") or "").lower()
            if label in {"page_footer", "page_header"}:
                continue
            page_number = self._node_page(node)
            blocks.append(
                ParsedBlock(
                    block_id=node.get("self_ref") or self._stable_block_id(Path("docling"), page_number, index, text[:80]),
                    block_index=len(blocks),
                    block_type=self._classify_node(node, text),
                    content=text,
                    page_number=page_number,
                    language=language,
                    bbox=self._node_bbox(node),
                    reading_order=len(blocks),
                    source="docling",
                    extra={"label": str(node.get("label") or node.get("type") or "")},
                )
            )

        # Add placeholder blocks for figures/pictures that have no caption/text.
        # For DOCX/PPTX exports, Docling often exposes embedded pictures without bbox.
        for fig_index, node in enumerate(figure_nodes):
            bbox = self._node_bbox(node)
            page_number = self._node_page(node)
            image_meta = node.get("image") if isinstance(node.get("image"), dict) else {}
            image_uri = image_meta.get("uri") if isinstance(image_meta, dict) else None
            block_id = node.get("self_ref") or self._stable_block_id(
                Path("docling"),
                page_number,
                10000 + fig_index,
                f"figure:{page_number}:{fig_index}:{str(node.get('self_ref') or '')}",
            )
            blocks.append(
                ParsedBlock(
                    block_id=block_id,
                    block_index=len(blocks),
                    block_type=BlockType.FIGURE.value,
                    content="",  # filled by FigureCaptioner in the pipeline
                    page_number=page_number,
                    language=language,
                    bbox=bbox,
                    reading_order=len(blocks),
                    source="docling",
                    extra={
                        "label": "figure",
                        "needs_captioning": True,
                        "embedded_image_uri": image_uri,
                        "embedded_image_mimetype": image_meta.get("mimetype") if isinstance(image_meta, dict) else None,
                        "embedded_image_size": image_meta.get("size") if isinstance(image_meta, dict) else None,
                        "docling_node_ref": node.get("self_ref"),
                    },
                )
            )

        return sorted(blocks, key=lambda b: (b.page_number, b.reading_order))

    def _docx_table_blocks(self, *, file_path: Path, language: str, reading_order_offset: int) -> list[ParsedBlock]:
        try:
            from docx import Document
        except ImportError:
            logger.debug("python-docx is unavailable; skipping DOCX table augmentation")
            return []
        try:
            document = Document(str(file_path))
        except Exception as exc:
            logger.debug("Could not inspect DOCX tables", extra={"path": str(file_path), "error": str(exc)})
            return []

        blocks: list[ParsedBlock] = []
        for table_index, table in enumerate(document.tables):
            rows = [
                [self._normalize_cell_text(cell.text) for cell in row.cells]
                for row in table.rows
            ]
            rows = [row for row in rows if any(cell for cell in row)]
            if len(rows) < 2 or max((len(row) for row in rows), default=0) < 2:
                continue
            markdown = self._rows_to_markdown_table(rows)
            if not markdown:
                continue
            blocks.append(
                ParsedBlock(
                    block_id=self._stable_block_id(file_path, 1, 20000 + table_index, markdown[:80]),
                    block_index=reading_order_offset + len(blocks),
                    block_type=BlockType.TABLE.value,
                    content=markdown,
                    page_number=1,
                    language=language,
                    reading_order=reading_order_offset + len(blocks),
                    source="python_docx",
                    extra={"label": "docx_table", "table_index": table_index},
                )
            )
        return blocks

    @staticmethod
    def _normalize_cell_text(text: str) -> str:
        return " ".join(text.split())

    @staticmethod
    def _rows_to_markdown_table(rows: list[list[str]]) -> str:
        width = max((len(row) for row in rows), default=0)
        if width < 2:
            return ""
        padded = [(row + [""] * width)[:width] for row in rows]
        header, body = padded[0], padded[1:]
        separator = ["---"] * width

        def render(row: list[str]) -> str:
            return "| " + " | ".join(cell.replace("|", "\\|") for cell in row) + " |"

        return "\n".join([render(header), render(separator), *(render(row) for row in body)])

    def _collect_figure_nodes(self, value: Any) -> list[dict[str, Any]]:
        """Collect figure/picture nodes that have no text.

        Docling may export embedded DOCX pictures without bbox, but the node still
        carries image metadata and can be captioned downstream.
        """
        nodes: list[dict[str, Any]] = []
        if isinstance(value, dict):
            label = str(value.get("label") or value.get("type") or "").lower()
            has_text = bool(self._node_text(value))
            has_image = isinstance(value.get("image"), dict) and bool(value["image"].get("uri"))
            if ("picture" in label or "figure" in label) and not has_text and (has_image or self._node_bbox(value) is not None):
                nodes.append(value)
            for child in value.values():
                nodes.extend(self._collect_figure_nodes(child))
        elif isinstance(value, list):
            for item in value:
                nodes.extend(self._collect_figure_nodes(item))
        return nodes

    def _collect_text_nodes(self, value: Any) -> list[dict[str, Any]]:
        nodes: list[dict[str, Any]] = []
        if isinstance(value, dict):
            if self._node_text(value):
                nodes.append(value)
            for child in value.values():
                nodes.extend(self._collect_text_nodes(child))
        elif isinstance(value, list):
            for item in value:
                nodes.extend(self._collect_text_nodes(item))
        return nodes

    @staticmethod
    def _node_text(node: dict[str, Any]) -> str:
        for key in ("text", "orig", "content"):
            value = node.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()
        return ""

    @staticmethod
    def _node_page(node: dict[str, Any]) -> int:
        prov = node.get("prov")
        if isinstance(prov, list) and prov:
            page_no = prov[0].get("page_no") if isinstance(prov[0], dict) else None
            if isinstance(page_no, int):
                return page_no
        page_no = node.get("page_no") or node.get("page")
        return int(page_no) if isinstance(page_no, int | float | str) and str(page_no).isdigit() else 1

    @staticmethod
    def _node_bbox(node: dict[str, Any]) -> BBox | None:
        prov = node.get("prov")
        candidate = None
        if isinstance(prov, list) and prov and isinstance(prov[0], dict):
            candidate = prov[0].get("bbox")
        candidate = candidate or node.get("bbox")
        if isinstance(candidate, dict):
            left = candidate.get("l", candidate.get("left", candidate.get("x1")))
            top = candidate.get("t", candidate.get("top", candidate.get("y1")))
            right = candidate.get("r", candidate.get("right", candidate.get("x2")))
            bottom = candidate.get("b", candidate.get("bottom", candidate.get("y2")))
            if all(value is not None for value in (left, top, right, bottom)):
                return BBox(x1=float(left), y1=float(top), x2=float(right), y2=float(bottom))
        return None

    @staticmethod
    def _classify_node(node: dict[str, Any], text: str) -> str:
        label = str(node.get("label") or node.get("type") or "").lower()
        if "table" in label:
            return BlockType.TABLE.value
        if "formula" in label or "equation" in label:
            return BlockType.EQUATION.value
        if "list" in label:
            return BlockType.LIST.value
        if "picture" in label or "figure" in label:
            return BlockType.FIGURE.value
        if "section_header" in label or "heading" in label or "title" in label:
            return BlockType.HEADING.value
        if text.strip().startswith(("-", "*", "\u2022")):
            return BlockType.LIST.value
        return BlockType.PARAGRAPH.value

    @staticmethod
    def _stable_block_id(file_path: Path, page_number: int, index: int, seed: str) -> str:
        return f"blk-{uuid5(NAMESPACE_URL, f'{file_path}:{page_number}:{index}:{seed}').hex[:12]}"
